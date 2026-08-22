"""
Script tạo file Word báo cáo đồ án AI Room Recommendation.
Chạy: python scripts/generate_report.py
Output: Bao_Cao_Do_An_AI_Room_Recommendation.docx
"""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


ROOT = Path(__file__).parent.parent
OUTPUT_PATH = ROOT / "Bao_Cao_Tien_Do_v2.docx"


# ─── Helper functions ─────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex: str):
    """Tô màu nền cho cell trong table."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_table(doc, headers: list[str], rows: list[list[str]], col_widths=None):
    """Tạo table với header có màu nền."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = text
        set_cell_shading(cell, "1F4E79")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = "Times New Roman"

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_heading_numbered(doc, text: str, level: int):
    """Thêm heading có đánh số."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(31, 78, 121)
    return heading


def add_body_text(doc, text: str, bold=False, italic=False, indent=False):
    """Thêm đoạn văn bản body."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text: str, level: int = 0):
    """Thêm bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    return p


# ─── Main report generation ──────────────────────────────────────────────────

def generate_report():
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    # ── Modify default styles ─────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.5

    # ══════════════════════════════════════════════════════════════════════════
    # TRANG BÌA
    # ══════════════════════════════════════════════════════════════════════════

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BÁO CÁO TIẾN ĐỘ ĐỒ ÁN")
    run.font.name = "Times New Roman"
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("HỆ THỐNG GỢI Ý PHÒNG TRỌ THÔNG MINH\nSỬ DỤNG TRÍ TUỆ NHÂN TẠO")
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("AI Room Recommendation System")
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.italic = True
    run.font.color.rgb = RGBColor(89, 89, 89)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Content-Based Filtering + LightGBM Re-ranking")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.italic = True
    run.font.color.rgb = RGBColor(89, 89, 89)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tiến độ hoàn thành: 70%")
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Năm 2026")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # MỤC LỤC
    # ══════════════════════════════════════════════════════════════════════════

    add_heading_numbered(doc, "MỤC LỤC", level=0)

    toc_items = [
        "1. Giới thiệu",
        "    1.1. Bối cảnh và động lực",
        "    1.2. Quá trình phát triển — Từ FAISS đến Pipeline 2 tầng",
        "    1.3. Mục tiêu đồ án",
        "2. Những gì đã thực hiện",
        "    2.1. Xây dựng Feature Extraction (phòng trọ → vector)",
        "    2.2. Xây dựng FAISS Index (tìm kiếm vector)",
        "    2.3. Xây dựng User Profile (hiểu sở thích người dùng)",
        "    2.4. Xây dựng LightGBM Re-ranker (tầng 2)",
        "3. Tiến độ và hướng phát triển tiếp theo",
        "    3.1. Tổng kết tiến độ hiện tại",
        "    3.2. Hướng phát triển tiếp theo",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHƯƠNG 1: GIỚI THIỆU
    # ══════════════════════════════════════════════════════════════════════════

    add_heading_numbered(doc, "1. GIỚI THIỆU", level=1)

    # ── 1.1 ───────────────────────────────────────────────────────────────────
    add_heading_numbered(doc, "1.1. Bối cảnh và động lực", level=2)

    add_body_text(doc,
        "Thị trường phòng trọ tại các thành phố lớn của Việt Nam như Hà Nội, "
        "TP. Hồ Chí Minh và Đà Nẵng ngày càng phát triển mạnh mẽ với hàng nghìn "
        "tin đăng mỗi ngày. Người thuê phòng, đặc biệt là sinh viên và người lao động "
        "mới đến thành phố, thường gặp khó khăn trong việc tìm được phòng phù hợp "
        "với nhu cầu cá nhân về giá cả, vị trí, diện tích và tiện ích.",
        indent=True
    )

    add_body_text(doc,
        "Các nền tảng cho thuê phòng hiện tại chủ yếu dựa vào bộ lọc thủ công "
        "(filter theo giá, quận, loại phòng), buộc người dùng phải tự khám phá và "
        "so sánh từng tin đăng. Phương pháp này tốn thời gian, không cá nhân hóa "
        "và dễ bỏ sót những phòng phù hợp nhưng không nằm trong phạm vi tìm kiếm "
        "ban đầu của người dùng.",
        indent=True
    )

    add_body_text(doc,
        "Xuất phát từ thực tế đó, đồ án này xây dựng một hệ thống gợi ý phòng trọ "
        "thông minh (AI Room Recommendation System) sử dụng các kỹ thuật học máy "
        "và tìm kiếm vector để tự động đề xuất những phòng trọ phù hợp nhất cho "
        "từng người dùng, dựa trên hành vi xem phòng và thông tin cá nhân của họ.",
        indent=True
    )

    # ── 1.2 ───────────────────────────────────────────────────────────────────
    add_heading_numbered(doc, "1.2. Quá trình phát triển — Từ FAISS đến Pipeline 2 tầng", level=2)

    add_body_text(doc,
        "Trong giai đoạn đầu phát triển, hệ thống chỉ sử dụng một mô hình duy nhất "
        "dựa trên FAISS (Facebook AI Similarity Search) — một thư viện tìm kiếm "
        "vector hiệu suất cao do Facebook AI Research phát triển. Cách tiếp cận ban đầu "
        "hoạt động theo nguyên lý Content-Based Filtering:",
        indent=True
    )

    add_bullet(doc, "Mỗi phòng trọ được biểu diễn dưới dạng một vector đặc trưng "
        "(feature vector) dựa trên các thuộc tính: giá, diện tích, quận/huyện, "
        "thành phố, loại phòng và tiện ích.")
    add_bullet(doc, "Sở thích của người dùng được tổng hợp thành một \"user profile vector\" "
        "từ lịch sử xem phòng (weighted average với time decay).")
    add_bullet(doc, "FAISS thực hiện tìm kiếm K-nearest neighbors trong không gian vector "
        "để tìm các phòng có cosine similarity cao nhất với profile của user.")

    add_body_text(doc,
        "Mô hình FAISS đơn tầng này cho kết quả khá tốt ở mức cơ bản — có thể tìm được "
        "các phòng tương tự về mặt nội dung. Tuy nhiên, qua quá trình kiểm thử và đánh giá, "
        "một số hạn chế rõ ràng đã được nhận ra:",
        indent=True
    )

    add_bullet(doc, "FAISS chỉ so sánh \"giống nhau\" về nội dung (content similarity), "
        "nhưng không nắm bắt được mức độ \"phù hợp\" thực sự giữa user và phòng. "
        "Ví dụ: hai phòng có vector tương tự nhau nhưng phòng có giá nằm trong ngân sách "
        "của user rõ ràng phù hợp hơn phòng vượt ngân sách.")
    add_bullet(doc, "Không có khả năng học các tương tác phức tạp giữa đặc trưng user "
        "và đặc trưng phòng, chẳng hạn: \"user này thường xem phòng ở Quận 1 với giá "
        "3-5 triệu\" — FAISS không mô hình hóa được mối quan hệ chéo này.")
    add_bullet(doc, "Thiếu các tín hiệu ngữ cảnh như khoảng cách địa lý thực tế, "
        "mức độ khớp tiện ích (amenity Jaccard), hay phòng có giá nằm trong phạm vi "
        "1 độ lệch chuẩn so với sở thích user hay không.")

    add_body_text(doc,
        "Nhận thấy những hạn chế trên, em đã quyết định bổ sung tầng thứ 2 vào pipeline "
        "— một mô hình LightGBM Re-ranker. Ý tưởng là sử dụng FAISS như một bộ lọc nhanh "
        "(retrieval stage) để thu hẹp từ hàng nghìn phòng xuống còn ~80 ứng viên tốt nhất, "
        "sau đó dùng LightGBM với ~32 features phong phú hơn để xếp hạng lại (re-ranking "
        "stage) và chọn ra top-K kết quả cuối cùng.",
        indent=True
    )

    add_body_text(doc,
        "Kiến trúc 2 tầng này kết hợp được ưu điểm của cả hai phương pháp: tốc độ tìm kiếm "
        "cực nhanh của FAISS (sub-millisecond trên hàng nghìn phòng) với khả năng học "
        "các interaction features phức tạp của LightGBM, mang lại kết quả gợi ý chính xác "
        "và cá nhân hóa hơn đáng kể.",
        indent=True
    )

    # ── 1.3 ───────────────────────────────────────────────────────────────────
    add_heading_numbered(doc, "1.3. Mục tiêu đồ án", level=2)

    add_bullet(doc, "Xây dựng hệ thống gợi ý phòng trọ sử dụng pipeline AI 2 tầng "
        "(FAISS retrieval + LightGBM re-ranking).")
    add_bullet(doc, "Xử lý bài toán cold start cho user mới chưa có lịch sử xem phòng "
        "bằng cách tận dụng thông tin địa chỉ đăng ký.")
    add_bullet(doc, "Cung cấp REST API hoàn chỉnh để tích hợp với frontend/mobile app.")
    add_bullet(doc, "Hỗ trợ dữ liệu phòng trọ tại 3 thành phố: Hà Nội, TP.HCM, Đà Nẵng.")
    add_bullet(doc, "Thiết kế modular, dễ mở rộng và sẵn sàng triển khai production.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHƯƠNG 2: NHỮNG GÌ ĐÃ THỰC HIỆN
    # ══════════════════════════════════════════════════════════════════════════

    add_heading_numbered(doc, "2. NHỮNG GÌ ĐÃ THỰC HIỆN", level=1)

    # ── 2.1 Feature Extraction ────────────────────────────────────────────────
    add_heading_numbered(doc, "2.1. Xây dựng Feature Extraction (phòng trọ → vector)", level=2)

    add_body_text(doc, "Đã làm gì:", bold=True)
    add_body_text(doc,
        "Em đã xây dựng module trích xuất đặc trưng (feature extraction) để chuyển "
        "mỗi phòng trọ thành một vector số có 49 chiều. Vector này chứa toàn bộ thông tin "
        "quan trọng của phòng: giá, diện tích, vị trí, loại phòng và các tiện ích.",
        indent=True
    )

    add_body_text(doc, "Cách làm:", bold=True)

    add_bullet(doc, "Giá và diện tích: Chuẩn hóa về khoảng [0, 1] bằng min-max scaling "
        "(giá: 500K–20M VND, diện tích: 10–100m²), sau đó nhân với trọng số "
        "(giá × 2.0, diện tích × 1.5) vì giá là yếu tố quan trọng nhất.")
    add_bullet(doc, "Thành phố, quận, loại phòng: Mã hóa one-hot. Ví dụ: \"Ho Chi Minh\" "
        "→ [0, 0, 1], \"Quận 1\" → [0, 0, ..., 1, ..., 0]. Tổng cộng 54 chiều.")
    add_bullet(doc, "Tiện ích (wifi, điều hòa, WC riêng...): Mã hóa binary — có tiện ích "
        "= 1, không có = 0. Tổng cộng 9 chiều.")

    add_body_text(doc, "Thiết kế config-driven:", bold=True)
    add_body_text(doc,
        "Toàn bộ cấu hình features (danh sách quận, loại phòng, tiện ích, trọng số...) "
        "được lưu trong file feature_config.json. Khi cần thêm quận mới hoặc tiện ích mới, "
        "chỉ cần cập nhật file JSON này rồi chạy lại indexer — không cần sửa code Python.",
        indent=True
    )

    # ── 2.2 FAISS Indexing ────────────────────────────────────────────────────
    add_heading_numbered(doc, "2.2. Xây dựng FAISS Index (tìm kiếm vector)", level=2)

    add_body_text(doc, "Đã làm gì:", bold=True)
    add_body_text(doc,
        "Em đã xây dựng module indexing sử dụng thư viện FAISS để lưu trữ tất cả "
        "vector phòng trọ và hỗ trợ tìm kiếm nearest neighbors cực nhanh. "
        "Đây là thành phần cốt lõi của tầng 1 trong pipeline.",
        indent=True
    )

    add_body_text(doc, "Cách làm:", bold=True)

    add_bullet(doc, "L2-normalize toàn bộ vector trước khi đưa vào index. "
        "Nhờ đó, inner product = cosine similarity — đo mức độ \"giống nhau\" giữa 2 phòng.")
    add_bullet(doc, "Tự động chọn loại index phù hợp: nếu dưới 5,000 phòng thì dùng "
        "IndexFlatIP (tìm kiếm chính xác 100%); nếu trên 5,000 phòng thì dùng "
        "IndexIVFFlat (tìm gần đúng, nhanh hơn 50-100 lần, recall vẫn ≥ 95%).")
    add_bullet(doc, "Lưu index xuống disk bằng cơ chế atomic swap: ghi vào file tạm "
        "trước, hoàn tất rồi mới rename — tránh lỗi nếu app đang đọc file cũ.")
    add_bullet(doc, "Có tính năng auto-tune nprobe: thử nhiều giá trị nprobe khác nhau, "
        "chọn giá trị nhỏ nhất đạt recall@10 ≥ 95% để cân bằng tốc độ và chính xác.")

    # ── 2.3 User Profile ──────────────────────────────────────────────────────
    add_heading_numbered(doc, "2.3. Xây dựng User Profile (hiểu sở thích người dùng)", level=2)

    add_body_text(doc, "Đã làm gì:", bold=True)
    add_body_text(doc,
        "Em đã xây dựng module tính toán vector sở thích của từng user (user profile), "
        "dùng để truy vấn FAISS và tìm phòng phù hợp. Module này kết hợp 2 nguồn "
        "thông tin: lịch sử xem phòng và địa chỉ đăng ký.",
        indent=True
    )

    add_body_text(doc, "Cách làm:", bold=True)

    add_bullet(doc, "History Profile: Lấy tối đa 50 lượt xem gần nhất, trích xuất "
        "vector của mỗi phòng đã xem, gán trọng số time decay (e^(-0.005 × số_giờ)). "
        "Phòng xem hôm qua có weight ≈ 0.89, xem 1 tuần trước ≈ 0.43, "
        "xem 1 tháng trước ≈ 0.03. Tính weighted average → history profile.")
    add_bullet(doc, "Address Profile: Với user mới chưa xem phòng nào, tạo một "
        "\"phòng ảo\" chỉ chứa thông tin city/district của user, các thuộc tính khác "
        "để bằng 0 → FAISS sẽ ưu tiên gợi ý phòng cùng khu vực.")
    add_bullet(doc, "Blend 2 nguồn: Dùng hệ số alpha tăng dần từ 0 → 1 theo số lượt xem. "
        "User mới (0 lượt) → 100% address; user xem 5 lượt → 50/50; "
        "user xem ≥ 10 lượt → 100% history. Chuyển đổi mượt mà, không bị giật.")

    # ── 2.4 LightGBM Re-ranker ────────────────────────────────────────────────
    add_heading_numbered(doc, "2.4. Xây dựng LightGBM Re-ranker (tầng 2)", level=2)

    add_body_text(doc, "Đã làm gì:", bold=True)
    add_body_text(doc,
        "Em đã bổ sung tầng 2 vào pipeline — mô hình LightGBM dùng để xếp hạng lại "
        "(re-rank) các phòng ứng viên từ FAISS. Tầng này tính ~32 features tương tác "
        "giữa user và phòng, dự đoán xác suất user \"thích\" phòng đó.",
        indent=True
    )

    add_body_text(doc, "Cách làm:", bold=True)

    add_bullet(doc, "Tầng 1 (FAISS) lấy 80 phòng ứng viên có cosine similarity cao nhất.")
    add_bullet(doc, "Với mỗi cặp (user, phòng ứng viên), tính 32 features gồm: "
        "FAISS score, thống kê lịch sử user (giá trung bình, diện tích trung bình...), "
        "chênh lệch giá/diện tích so với sở thích, có cùng quận/thành phố không, "
        "Jaccard similarity tiện ích, khoảng cách địa lý (km)...")
    add_bullet(doc, "LightGBM (binary classification) dự đoán xác suất user sẽ xem/thích "
        "phòng → sort theo xác suất giảm dần → trả top-K.")
    add_bullet(doc, "Chiến lược bypass: nếu user mới xem < 5 phòng → bỏ qua LightGBM, "
        "giữ nguyên thứ tự FAISS. Lý do: khi data quá ít, các interaction features "
        "gần như toàn số 0, LightGBM không thể ra quyết định tốt.")
    add_bullet(doc, "Huấn luyện: Positive samples = cặp (user, phòng đã xem); "
        "Negative samples = phòng trong top-K FAISS nhưng user không xem (hard negatives).")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHƯƠNG 3: TIẾN ĐỘ VÀ HƯỚNG PHÁT TRIỂN
    # ══════════════════════════════════════════════════════════════════════════

    add_heading_numbered(doc, "3. TIẾN ĐỘ VÀ HƯỚNG PHÁT TRIỂN TIẾP THEO", level=1)

    add_heading_numbered(doc, "3.1. Tổng kết tiến độ hiện tại", level=2)

    add_body_text(doc,
        "Dưới đây là bảng tổng kết các đầu việc chính của đồ án, bao gồm những phần "
        "đã hoàn thành và những phần đang trong kế hoạch thực hiện tiếp theo:",
        indent=True
    )

    add_formatted_table(doc,
        headers=["STT", "Công việc", "Trạng thái", "Ghi chú"],
        rows=[
            ["1", "Feature Extraction (phòng → vector 49 chiều)", "Hoàn thành ✓",
             "Config-driven, dễ mở rộng"],
            ["2", "FAISS Indexing (tìm kiếm vector)", "Hoàn thành ✓",
             "Auto chọn index type, auto-tune nprobe"],
            ["3", "User Profile Building", "Hoàn thành ✓",
             "History + address blend, time decay"],
            ["4", "LightGBM Re-ranker (tầng 2)", "Hoàn thành ✓",
             "32 features, bypass khi data ít"],
            ["5", "Xây dựng API Server", "Chưa thực hiện",
             "REST API để frontend gọi"],
            ["6", "Xử lý Cold Start", "Chưa thực hiện",
             "User mới chưa có lịch sử"],
            ["7", "Mở rộng dữ liệu", "Chưa thực hiện",
             "Lượng data hiện tại chưa đủ lớn"],
            ["8", "Đánh giá mô hình", "Chưa thực hiện",
             "Metrics: Precision@K, Recall@K, NDCG"],
        ],
        col_widths=[1, 6.5, 3.5, 5]
    )

    add_heading_numbered(doc, "3.2. Hướng phát triển tiếp theo", level=2)

    add_body_text(doc,
        "Những phần dự kiến sẽ tiếp tục thực hiện để hoàn thiện đồ án:",
        indent=True
    )

    # ── 3.2.1 API Server ──────────────────────────────────────────────────────
    add_body_text(doc, "a) Xây dựng API Server", bold=True)
    add_body_text(doc,
        "Xây dựng REST API server bằng FastAPI để cung cấp các endpoints cho frontend/mobile app "
        "có thể truy vấn và sử dụng hệ thống gợi ý. Dự kiến các endpoints chính:",
        indent=True
    )
    add_bullet(doc, "API lấy danh sách phòng, chi tiết phòng (có filter theo giá, quận, loại phòng).")
    add_bullet(doc, "API gợi ý phòng cho user (GET /recommendations/{user_id}) — endpoint chính "
        "gọi vào pipeline 2 tầng FAISS + LightGBM.")
    add_bullet(doc, "API ghi nhận sự kiện xem phòng — cập nhật lịch sử user để cải thiện gợi ý.")
    add_bullet(doc, "API admin: rebuild FAISS index và retrain LightGBM (chạy nền, không block service).")

    # ── 3.2.2 Cold Start ──────────────────────────────────────────────────────
    add_body_text(doc, "b) Xử lý Cold Start", bold=True)
    add_body_text(doc,
        "Xây dựng chiến lược xử lý bài toán cold start — khi user mới đăng ký chưa có "
        "lịch sử xem phòng. Dự kiến xử lý theo 3 mức:",
        indent=True
    )
    add_bullet(doc, "User mới có địa chỉ đăng ký: tạo address profile từ city/district "
        "→ FAISS tìm phòng cùng khu vực mà không cần lịch sử.")
    add_bullet(doc, "User xem được vài phòng (< 10 lượt): blend giữa address profile "
        "và history profile, bypass LightGBM vì data chưa đủ tin cậy.")
    add_bullet(doc, "User không có thông tin gì: popular fallback — gợi ý random phòng "
        "còn trống trong hệ thống.")

    # ── 3.2.3 Dữ liệu ────────────────────────────────────────────────────────
    add_body_text(doc, "c) Mở rộng và bổ sung dữ liệu", bold=True)
    add_body_text(doc,
        "Hiện tại, lượng dữ liệu sử dụng trong hệ thống chưa đủ lớn để đánh giá "
        "chính xác hiệu quả của mô hình. Cụ thể, bộ dữ liệu hiện có bao gồm "
        "rooms.json (~5 MB), users.json (~31 KB) và view_history.json (~769 KB), "
        "phạm vi 3 thành phố với 47 quận/huyện. Tuy nhiên, số lượng user và lịch sử "
        "tương tác còn hạn chế, dẫn đến việc đánh giá metrics (Precision, Recall, NDCG) "
        "có thể chưa phản ánh đúng năng lực thực tế của mô hình.",
        indent=True
    )
    add_body_text(doc, "Hướng mở rộng dữ liệu:", bold=True)
    add_bullet(doc, "Thu thập thêm dữ liệu phòng trọ thực tế từ các nền tảng cho thuê phòng "
        "(crawl/API) để tăng số lượng phòng và đa dạng hóa dữ liệu.")
    add_bullet(doc, "Tăng số lượng user giả lập và lịch sử xem phòng để mô hình LightGBM "
        "có đủ dữ liệu huấn luyện, đặc biệt là các interaction features.")
    add_bullet(doc, "Bổ sung thêm implicit feedback: thời gian xem chi tiết phòng, "
        "số lần quay lại xem, lưu phòng yêu thích — giúp mô hình hiểu sâu hơn "
        "về sở thích thực sự của user.")
    add_bullet(doc, "Cần dữ liệu đủ lớn để đánh giá kết quả chính xác hơn, "
        "đặc biệt là so sánh hiệu quả giữa FAISS đơn tầng và pipeline 2 tầng "
        "trên các phân khúc user khác nhau (user mới, user cũ, user nhiều/ít lịch sử).")

    # ── 3.2.4 Đánh giá ────────────────────────────────────────────────────────
    add_body_text(doc, "d) Đánh giá định lượng hiệu quả mô hình", bold=True)
    add_body_text(doc,
        "Sau khi có đủ dữ liệu, tiến hành đánh giá mô hình bằng các metrics chuẩn:",
        indent=True
    )
    add_bullet(doc, "Precision@K, Recall@K: đo tỉ lệ gợi ý đúng trong top-K kết quả.")
    add_bullet(doc, "NDCG (Normalized Discounted Cumulative Gain): đánh giá chất lượng "
        "xếp hạng — phòng phù hợp nhất có được xếp đầu tiên không.")
    add_bullet(doc, "So sánh A/B: FAISS đơn tầng vs Pipeline 2 tầng (FAISS + LightGBM) "
        "→ chứng minh tầng 2 thực sự cải thiện chất lượng gợi ý.")
    add_bullet(doc, "Phân tích theo phân khúc: hiệu quả trên user mới (cold start) "
        "vs user có nhiều lịch sử xem.")

    # ── 3.2.5 Khác ────────────────────────────────────────────────────────────
    add_body_text(doc, "e) Các công việc khác", bold=True)
    add_bullet(doc, "Tích hợp frontend: kết nối API với giao diện web/mobile để "
        "demo trực quan.")
    add_bullet(doc, "Tối ưu hiệu năng: cache user profile, batch prediction, "
        "benchmark thời gian response.")
    add_bullet(doc, "Hoàn thiện báo cáo cuối: bổ sung kết quả đánh giá, "
        "phân tích ưu/nhược điểm và hướng phát triển dài hạn.")

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(str(OUTPUT_PATH))
    print(f"Done! File: {OUTPUT_PATH.resolve()}")


if __name__ == "__main__":
    generate_report()
